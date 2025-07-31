# 📄 utils/extractor.py
import fitz  # PyMuPDF
import pdfplumber

def extract_text_from_pdf(pdf_path):
    text = ""
    # Try pdfplumber first
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        print("pdfplumber failed:", e)

    # Fallback to PyMuPDF if pdfplumber fails or returns nothing
    if not text.strip():
        try:
            doc = fitz.open(pdf_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception as e:
            print("PyMuPDF fallback failed:", e)

    return text.strip()


def extract_text_from_docx(docx_path):
    from docx import Document
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text(file_path: str) -> str:
    # Try using pdfplumber
    if file_path.lower().endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(page.extract_text() or '' for page in pdf.pages)
                if text.strip():
                    return text
        except Exception as e:
            print(f"[pdfplumber failed] {e}")

        # Fall back to PyMuPDF
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = "\n".join(page.get_text() for page in doc)
            return text.strip()
        except Exception as e:
            print(f"[fitz failed] {e}")
            return ""
    elif file_path.lower().endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_path.lower().endswith(".docx"):
        import docx
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""
